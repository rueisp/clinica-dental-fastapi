from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generar_historia_clinica_word(paciente, odontologo):
    doc = Document()
    
    # 1. AJUSTE DE MÁRGENES ESTRECHOS
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- ENCABEZADO ---
    titulo = doc.add_heading('HISTORIA CLÍNICA ODONTOLÓGICA', 1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(10)

    # Función auxiliar para añadir filas a las tablas de datos
    def agregar_fila_datos(tabla, etiqueta1, valor1, etiqueta2, valor2):
        row_cells = tabla.add_row().cells
        # Col 1
        p1 = row_cells[0].paragraphs[0]
        p1.add_run(f"{etiqueta1}: ").bold = True
        p1.add_run(str(valor1 or ""))
        p1.paragraph_format.space_after = Pt(1)
        # Col 2
        p2 = row_cells[1].paragraphs[0]
        p2.add_run(f"{etiqueta2}: ").bold = True
        p2.add_run(str(valor2 or ""))
        p2.paragraph_format.space_after = Pt(1)
        # Tamaño de fuente para toda la fila
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    # --- SECCIÓN 1: DATOS PERSONALES ---
    doc.add_heading('1. Información Personal', level=2)
    tabla_pers = doc.add_table(rows=0, cols=2)
    tabla_pers.style = 'Table Grid' # <--- Agrega esto

    fecha_nac = paciente.fecha_nacimiento.strftime('%d/%m/%Y') if paciente.fecha_nacimiento else ""

    agregar_fila_datos(tabla_pers, "Paciente", f"{paciente.nombres} {paciente.apellidos}", "Identificación", f"{paciente.tipo_documento or ''} {paciente.documento or ''}")
    agregar_fila_datos(tabla_pers, "F. Nacimiento", fecha_nac, "Edad/Sexo", f"{paciente.edad or ''} años / {paciente.sexo or ''}")
    agregar_fila_datos(tabla_pers, "Ocupación", paciente.ocupacion, "Teléfono", paciente.telefono)
    agregar_fila_datos(tabla_pers, "Email", paciente.email, "Dirección", paciente.direccion)


    # --- SECCIÓN 2: ANAMNESIS Y ANTECEDENTES ---
    h2 = doc.add_heading('2. Anamnesis y Antecedentes', level=2)
    h2.paragraph_format.space_before = Pt(8)
    
    # Usamos una tabla de una sola columna para textos largos (Motivo, Enfermedad, etc.)
    tabla_clin = doc.add_table(rows=0, cols=1)
    tabla_clin.style = 'Table Grid'

    def agregar_bloque_texto(titulo_bloque, contenido):
        row = tabla_clin.add_row().cells[0]
        p = row.paragraphs[0]
        p.add_run(f"{titulo_bloque}: ").bold = True
        p.add_run(str(contenido or "No registra"))
        p.runs[0].font.size = Pt(8.5)
        p.runs[1].font.size = Pt(8.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    agregar_bloque_texto("Motivo de Consulta", paciente.motivo_consulta)
    agregar_bloque_texto("Enfermedad Actual", paciente.enfermedad_actual)
    agregar_bloque_texto("Alergias", paciente.alergias)
    agregar_bloque_texto("Hábitos", paciente.habitos)
    agregar_bloque_texto("Cepillado Dental", paciente.cepillado_dental)
    agregar_bloque_texto("Observaciones Generales", paciente.observaciones)

    # --- SECCIÓN 3: EVOLUCIONES ---
    h3 = doc.add_heading('3. Evoluciones Clínicas', level=2)
    h3.paragraph_format.space_before = Pt(8)

    evoluciones_ordenadas = sorted(paciente.evoluciones, key=lambda x: x.fecha, reverse=True)

    if not evoluciones_ordenadas:
        p = doc.add_paragraph("No se registran evoluciones.")
        p.font.size = Pt(8.5)
    else:
        for ev in evoluciones_ordenadas:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            
            run_fecha = p.add_run(f"{ev.fecha.strftime('%d/%m/%Y %H:%M')} - ")
            run_fecha.bold = True
            run_fecha.font.size = Pt(8.5)
            
            run_desc = p.add_run(ev.descripcion)
            run_desc.font.size = Pt(8.5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- FIRMA ---
    doc.add_paragraph("\n")
    firma_p = doc.add_paragraph()
    firma_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_firma = firma_p.add_run(f"__________________________\nDr/Dra. {odontologo.nombres or odontologo.username}\nOdontólogo Responsable")
    run_firma.font.size = Pt(8)

    # Ajustar estilos de los Headings para que sean más pequeños
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 1'):
            for run in paragraph.runs: run.font.size = Pt(14)
        if paragraph.style.name.startswith('Heading 2'):
            for run in paragraph.runs: run.font.size = Pt(10)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer